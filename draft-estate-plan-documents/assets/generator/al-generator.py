#!/usr/bin/env python3
"""Estate planning document generator.

This tool renders Word templates with structured client data, reusable clauses,
and state-specific overrides. It is intentionally conservative: unresolved
placeholders, missing required data, and unapproved clauses are reported instead
of being silently ignored.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, StrictUndefined, TemplateError, Undefined


ROOT = Path(__file__).resolve().parent
PLACEHOLDER_RE = re.compile(r"{{\s*([^{}]+?)\s*}}")
BLOCK_RE = re.compile(r"{%\s*(.*?)\s*%}")
UNRESOLVED_RE = re.compile(r"({{.*?}}|{%.*?%}|\[[A-Z][A-Z0-9 _/-]*\])")
BODY_FONT_NAME = "Garamond"
BODY_FONT_SIZE_PT = 12
FOOTER_FONT_SIZE_PT = 10


class GeneratorError(Exception):
    """Raised when generation cannot safely continue."""


@dataclass
class GenerationReport:
    data_file: str
    state: str
    package: str | None = None
    template: str | None = None
    output_files: list[str] = field(default_factory=list)
    included_clauses: list[dict[str, Any]] = field(default_factory=list)
    skipped_clauses: list[dict[str, Any]] = field(default_factory=list)
    missing_fields: list[str] = field(default_factory=list)
    unresolved_markers: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "data_file": self.data_file,
            "state": self.state,
            "package": self.package,
            "template": self.template,
            "output_files": self.output_files,
            "included_clauses": self.included_clauses,
            "skipped_clauses": self.skipped_clauses,
            "missing_fields": sorted(set(self.missing_fields)),
            "unresolved_markers": sorted(set(self.unresolved_markers)),
            "warnings": self.warnings,
            "errors": self.errors,
        }


def load_json(path: Path) -> Any:
    try:
        with path.open("r", encoding="utf-8") as handle:
            return json.load(handle)
    except FileNotFoundError as exc:
        raise GeneratorError(f"File not found: {path}") from exc
    except json.JSONDecodeError as exc:
        raise GeneratorError(f"Invalid JSON in {path}: {exc}") from exc


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
        handle.write("\n")


def get_path(data: Any, dotted_path: str) -> Any:
    current = data
    for raw_part in dotted_path.split("."):
        part = raw_part.removesuffix("[]")
        if isinstance(current, list):
            if not current:
                return None
            current = [item.get(part) if isinstance(item, dict) else None for item in current]
            continue
        if not isinstance(current, dict) or part not in current:
            return None
        current = current[part]
    return current


def is_present(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, (list, dict)):
        return bool(value)
    return True


def format_name(value: Any) -> str:
    if value is None:
        return ""
    return str(value).upper()


def build_environment(fail_fast: bool) -> Environment:
    env = Environment(undefined=StrictUndefined if fail_fast else Undefined, autoescape=False)
    env.filters["name"] = format_name
    return env


def missing_required_fields(data: dict[str, Any], required_fields: list[str]) -> list[str]:
    return [field for field in required_fields if not is_present(get_path(data, field))]


def load_state_rules(state: str) -> dict[str, Any]:
    rules_path = ROOT / "states" / state / "rules.json"
    if rules_path.exists():
        return load_json(rules_path)
    raise GeneratorError(f"Unsupported state: {state}. This project is currently configured for MA only.")


def compute_conditions(data: dict[str, Any], state: str) -> dict[str, Any]:
    spouse = data.get("spouse") or {}
    children = data.get("children") or []
    gifts = data.get("specific_gifts") or []
    charitable = data.get("charitable_gifts") or []
    planning = data.get("planning") or {}
    family = data.get("family") or {}
    assets = data.get("assets") or {}
    tax = data.get("tax") or {}
    guardians = data.get("guardians") or {}
    fiduciaries = data.get("fiduciaries") or {}
    agent = (fiduciaries.get("agent") or {}) if isinstance(fiduciaries, dict) else {}
    state_rules = load_state_rules(state)

    has_minor_children = any(bool(child.get("minor")) for child in children if isinstance(child, dict))
    has_spouse = bool(spouse.get("exists")) or is_present(spouse.get("full_name"))
    has_co_poa = is_present(get_path(data, "fiduciaries.agent.co_primary.full_name"))
    poa_action = agent.get("co_agents_action", "separate")

    return {
        "state": {"code": state, **state_rules},
        "has_spouse": has_spouse,
        "has_children": bool(children),
        "has_minor_children": has_minor_children,
        "has_blended_family": bool(family.get("blended_family")),
        "has_co_poa": has_co_poa,
        "has_solo_poa": not has_co_poa,
        "co_poa_joint": has_co_poa and poa_action == "joint",
        "co_poa_separate": has_co_poa and poa_action != "joint",
        "aif_has_spouse": bool(agent.get("primary_has_spouse")),
        "uses_revocable_trust": bool(planning.get("uses_revocable_trust")),
        "requires_guardian_clause": has_minor_children,
        "has_specific_gifts": bool(gifts),
        "has_charitable_gifts": bool(charitable),
        "has_excluded_guardians": bool(guardians.get("excluded")),
        "owns_real_property": bool(assets.get("real_property")),
        "tax_sensitive_estate": bool(tax.get("sensitive_estate")),
        "state_requires_self_proving_affidavit": bool(state_rules.get("requires_self_proving_affidavit")),
    }


class ClauseStore:
    def __init__(self, state: str, report: GenerationReport, allow_draft: bool = False) -> None:
        self.state = state
        self.report = report
        self.allow_draft = allow_draft
        self.base_clauses = self._load_clause_dir(ROOT / "clauses")
        self.state_clauses = self._load_clause_dir(ROOT / "states" / state / "clauses")

    def _load_clause_dir(self, directory: Path) -> dict[str, dict[str, Any]]:
        clauses: dict[str, dict[str, Any]] = {}
        if not directory.exists():
            return clauses
        for path in sorted(directory.rglob("*.json")):
            payload = load_json(path)
            clause_id = payload.get("id")
            if not clause_id:
                self.report.warnings.append(f"Clause file has no id and was ignored: {path}")
                continue
            payload["_source_file"] = str(path.relative_to(ROOT))
            clauses[clause_id] = payload
        return clauses

    def get(self, clause_id: str) -> dict[str, Any]:
        clause = self.state_clauses.get(clause_id) or self.base_clauses.get(clause_id)
        if not clause:
            raise GeneratorError(f"Clause not found: {clause_id}")
        return clause

    def include(self, clause_id: str, context: dict[str, Any], state: str | None = None) -> str:
        lookup_state = state or self.state
        if lookup_state != self.state:
            self.report.warnings.append(
                f"Requested clause {clause_id} for state {lookup_state}; active state is {self.state}."
            )
        clause = self.get(clause_id)
        condition = clause.get("condition")
        if condition and not bool(context.get(condition)):
            self.report.skipped_clauses.append(
                {"id": clause_id, "reason": f"condition false: {condition}", "source": clause.get("_source_file")}
            )
            return ""

        review_status = clause.get("review_status", "draft")
        if review_status != "approved" and not self.allow_draft:
            raise GeneratorError(
                f"Clause {clause_id} is marked {review_status!r}. Use --allow-draft to generate draft documents."
            )

        missing = missing_required_fields(context, clause.get("required_fields", []))
        if missing:
            self.report.missing_fields.extend(missing)
            raise GeneratorError(f"Clause {clause_id} requires missing field(s): {', '.join(missing)}")

        self.report.included_clauses.append(
            {
                "id": clause_id,
                "title": clause.get("title"),
                "jurisdiction": clause.get("jurisdiction"),
                "source": clause.get("_source_file"),
                "review_status": review_status,
            }
        )
        return render_string(clause.get("text", ""), context, self, fail_fast=True)


def render_string(source: str, context: dict[str, Any], clauses: ClauseStore, fail_fast: bool) -> str:
    env = build_environment(fail_fast)

    def include_clause(clause_id: str, state: str | None = None) -> str:
        return clauses.include(clause_id, context, state=state)

    env.globals["include_clause"] = include_clause
    try:
        return env.from_string(source).render(**context)
    except TemplateError as exc:
        raise GeneratorError(str(exc)) from exc


def context_for(data: dict[str, Any], state: str) -> dict[str, Any]:
    context = dict(data)
    context.update(compute_conditions(data, state))
    client = context.setdefault("client", {})
    if isinstance(client, dict):
        client.setdefault("state", state)
    trust = context.setdefault("trust", {})
    if isinstance(trust, dict):
        trust["display_date"] = trust.get("date") or "________________, 20___"
    execution = context.setdefault("execution", {})
    if isinstance(execution, dict):
        signing_scheduled = bool(execution.get("signing_scheduled"))
        context["signing_scheduled"] = signing_scheduled
        execution["display_doc_date"] = execution.get("doc_date") if signing_scheduled else "________________, 20__"
        execution["display_ordinal_doc_date"] = (
            execution.get("ordinal_doc_date") if signing_scheduled else "_____ day of _________________, 20___"
        )
        signing_county = execution.get("signing_county") or client.get("signing_county")
        execution["display_signing_county"] = signing_county if signing_scheduled else "__________________________"
        execution["display_notary_commission"] = (
            execution.get("notary_commission") if signing_scheduled else "_________________________"
        )
    return context


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def apply_run_font(paragraph: Any, size_pt: int) -> None:
    for run in paragraph.runs:
        run.font.name = BODY_FONT_NAME
        run.font.size = Pt(size_pt)


def apply_document_formatting(path: Path) -> None:
    document = Document(path)

    for style_name, size_pt in (("Normal", BODY_FONT_SIZE_PT), ("Footer", FOOTER_FONT_SIZE_PT)):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = BODY_FONT_NAME
            style.font.size = Pt(size_pt)

    for paragraph in document.paragraphs:
        apply_run_font(paragraph, BODY_FONT_SIZE_PT)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_run_font(paragraph, BODY_FONT_SIZE_PT)

    for section in document.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in footer.paragraphs:
                apply_run_font(paragraph, FOOTER_FONT_SIZE_PT)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            apply_run_font(paragraph, FOOTER_FONT_SIZE_PT)

    document.save(path)


def scan_template(path: Path) -> dict[str, list[str]]:
    text = extract_docx_text(path)
    placeholders = sorted(set(match.group(1).strip() for match in PLACEHOLDER_RE.finditer(text)))
    blocks = sorted(set(match.group(1).strip() for match in BLOCK_RE.finditer(text)))
    return {"placeholders": placeholders, "blocks": blocks}


def render_docx(
    template_path: Path,
    output_path: Path,
    context: dict[str, Any],
    clauses: ClauseStore,
    report: GenerationReport,
) -> None:
    if not template_path.exists():
        raise GeneratorError(f"Template not found: {template_path}")

    def include_clause(clause_id: str, state: str | None = None) -> str:
        return clauses.include(clause_id, context, state=state)

    doc = DocxTemplate(str(template_path))
    jinja_env = build_environment(fail_fast=True)
    try:
        doc.render({**context, "include_clause": include_clause}, jinja_env=jinja_env)
    except TemplateError as exc:
        raise GeneratorError(str(exc)) from exc
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    apply_document_formatting(output_path)

    unresolved = UNRESOLVED_RE.findall(extract_docx_text(output_path))
    if unresolved:
        report.unresolved_markers.extend(unresolved)
    report.output_files.append(str(output_path.relative_to(ROOT)))


def validate_data(context: dict[str, Any], required_fields: list[str], report: GenerationReport) -> None:
    report.missing_fields.extend(missing_required_fields(context, required_fields))
    execution = context.get("execution") or {}
    if execution.get("signing_scheduled"):
        report.missing_fields.extend(
            missing_required_fields(
                context,
                [
                    "execution.doc_date",
                    "execution.ordinal_doc_date",
                    "execution.display_signing_county",
                    "execution.notary_commission",
                ],
            )
        )


def package_config(package_name: str) -> dict[str, Any]:
    path = ROOT / "packages" / f"{package_name}.json"
    return load_json(path)


def run_validate(args: argparse.Namespace) -> int:
    data_path = Path(args.data)
    data = load_json(data_path)
    state = args.state or get_path(data, "client.state") or "base"
    context = context_for(data, state)
    report = GenerationReport(data_file=str(data_path), state=state, template=args.template)
    config_required = load_json(ROOT / "schemas" / "required-fields.json")
    validate_data(context, config_required.get("required_fields", []), report)

    if args.template:
        template_path = Path(args.template)
        if not template_path.is_absolute():
            template_path = ROOT / template_path
        scan = scan_template(template_path)
        report.warnings.append(f"Template placeholders found: {', '.join(scan['placeholders']) or 'none'}")
        report.warnings.append(f"Template logic tags found: {', '.join(scan['blocks']) or 'none'}")

    print(json.dumps(report.as_dict(), indent=2))
    return 1 if report.errors or report.missing_fields else 0


def build_generation_report(args: argparse.Namespace, render: bool) -> GenerationReport:
    data_path = Path(args.data)
    data = load_json(data_path)
    state = args.state or get_path(data, "client.state") or "base"
    package_name = args.package
    context = context_for(data, state)
    report = GenerationReport(data_file=str(data_path), state=state, package=package_name)
    config = package_config(package_name)
    validate_data(context, config.get("required_fields", []), report)
    if report.missing_fields:
        return report

    clauses = ClauseStore(state=state, report=report, allow_draft=args.allow_draft)
    out_dir = Path(getattr(args, "out", None) or "output")
    if not out_dir.is_absolute():
        out_dir = ROOT / out_dir

    for document in config.get("documents", []):
        template_path = ROOT / document["template"]
        output_name = document.get("output", template_path.name)
        output_path = out_dir / output_name
        try:
            if render:
                render_docx(template_path, output_path, context, clauses, report)
            else:
                template_text = extract_docx_text(template_path)
                render_string(template_text, context, clauses, fail_fast=True)
        except GeneratorError as exc:
            report.errors.append(f"{document.get('id', template_path.name)}: {exc}")
            if not render:
                continue
            break

    return report


def run_generate(args: argparse.Namespace) -> int:
    report = build_generation_report(args, render=True)
    out_dir = Path(args.out)
    if not out_dir.is_absolute():
        out_dir = ROOT / out_dir
    write_json(out_dir / "generation-report.json", report.as_dict())
    print(json.dumps(report.as_dict(), indent=2))
    return 1 if report.errors or report.missing_fields or report.unresolved_markers else 0


def run_report(args: argparse.Namespace) -> int:
    report = build_generation_report(args, render=False)
    print(json.dumps(report.as_dict(), indent=2))
    return 1 if report.errors or report.missing_fields else 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Estate planning Word document generator")
    subparsers = parser.add_subparsers(dest="command", required=True)

    validate = subparsers.add_parser("validate", help="Validate intake data and inspect a template")
    validate.add_argument("--data", required=True, help="Path to client JSON data")
    validate.add_argument("--template", help="Path to a .docx template to inspect")
    validate.add_argument("--state", help="State code override, such as MA")
    validate.set_defaults(func=run_validate)

    generate = subparsers.add_parser("generate", help="Generate a document package")
    generate.add_argument("--data", required=True, help="Path to client JSON data")
    generate.add_argument("--package", required=True, help="Package name, such as core-estate")
    generate.add_argument("--state", required=True, help="State code, such as MA")
    generate.add_argument("--out", required=True, help="Output directory")
    generate.add_argument("--allow-draft", action="store_true", help="Allow draft/reviewed clauses")
    generate.set_defaults(func=run_generate)

    report = subparsers.add_parser("report", help="Preview clause selection without writing documents")
    report.add_argument("--data", required=True, help="Path to client JSON data")
    report.add_argument("--package", required=True, help="Package name, such as core-estate")
    report.add_argument("--state", required=True, help="State code, such as MA")
    report.add_argument("--allow-draft", action="store_true", help="Allow draft/reviewed clauses")
    report.set_defaults(func=run_report)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return args.func(args)
    except GeneratorError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
