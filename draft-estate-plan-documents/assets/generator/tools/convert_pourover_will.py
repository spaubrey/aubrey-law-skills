#!/usr/bin/env python3
"""Convert the provided pour-over will into a generator-ready template."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source-forms" / "Pourover_Will.docx"
TARGET = ROOT / "templates" / "pourover-will.docx"


REPLACEMENTS = {
    "[CLIENT]": "{{ client.full_name | name }}",
    "[CITY]": "{{ client.city }}",
    "[COUNTY]": "{{ client.county }}",
    "[SPOUSE]": "{{ spouse.full_name | name }}",
    "[CHILD FULL NAME]": "{% for child in children %}{{ child.full_name | name }}{% if not loop.last %}, {% endif %}{% endfor %}",
    "[TRUST NAME]": "{{ trust.name | name }}",
    "[INITIAL PERSONAL REPRESENTATIVE]": "{{ fiduciaries.executor.primary.full_name | name }}",
    "[CO-INITIAL PERSONAL REPRESENTATIVE]": "{{ fiduciaries.executor.co_primary.full_name | name }}",
    "[SUCCESSOR PERSONAL REPRESENTATIVE]": "{{ fiduciaries.executor.successor.full_name | name }}",
    "[CO-SUCCESSOR PERSONAL REPRESENTATIVE]": "{{ fiduciaries.executor.co_successor.full_name | name }}",
    "[INITIAL GUARDIAN]": "{{ guardians.primary.full_name | name }}",
    "[SUCCESSOR GUARDIAN]": "{{ guardians.successor.full_name | name }}",
    "[SECOND SUCCESSOR GUARDIAN]": "{{ guardians.second_successor.full_name | name }}",
}


PARAGRAPH_REPLACEMENTS = {
    "NOMINATION OF GUARDIAN":
        "{% if requires_guardian_clause %}NOMINATION OF GUARDIAN{% endif %}",
    "I appoint as guardian of the person and the property of any minor child of mine, the child's other parent, if living, or, if that parent fails to qualify or ceases to act, I appoint [INITIAL GUARDIAN] as guardian of the person and the property of any such minor child of mine. If [INITIAL GUARDIAN] fails to qualify or ceases to act, I appoint [SUCCESSOR GUARDIAN], and if [SUCCESSOR GUARDIAN] fails to qualify or ceases to act, I appoint [SECOND SUCCESSOR GUARDIAN], as guardian of the person and the property of any such minor child of mine.  No bond or other security shall be required of any such person.  As used in this Article the term \"guardian of the property\" means the fiduciary (appointed by a court of competent jurisdiction or by other lawful means) responsible for the property of an individual and includes but is not limited to any similar terms such as conservator.":
        "{% if requires_guardian_clause %}I appoint as guardian of the person and the property of any minor child of mine, the child's other parent, if living, or, if that parent fails to qualify or ceases to act, I appoint {{ guardians.primary.full_name | name }} as guardian of the person and the property of any such minor child of mine. If {{ guardians.primary.full_name | name }} fails to qualify or ceases to act, I appoint {{ guardians.successor.full_name | name }}, and if {{ guardians.successor.full_name | name }} fails to qualify or ceases to act, I appoint {{ guardians.second_successor.full_name | name }}, as guardian of the person and the property of any such minor child of mine. No bond or other security shall be required of any such person. As used in this Article the term \"guardian of the property\" means the fiduciary (appointed by a court of competent jurisdiction or by other lawful means) responsible for the property of an individual and includes but is not limited to any similar terms such as conservator.{% endif %}",
    "I am married to [SPOUSE].  Any reference in this document to my spouse is a reference to [SPOUSE].":
        "{% if has_spouse %}I am married to {{ spouse.full_name | name }}. Any reference in this document to my spouse is a reference to {{ spouse.full_name | name }}.{% endif %}",
    "References to my children and descendants in this Last Will and Testament shall refer only to [CHILD FULL NAME], the descendants of my aforesaid child, any child of mine born or adopted after the date of this Last Will and Testament, and the descendants of any such after-born or after adopted child. A person legally adopted prior to attaining eighteen (18) years of age shall not be differentiated from blood descendants for any purpose, but a person who is adopted after attaining eighteen (18) years of age shall not be deemed a child or descendant for all purposes hereunder.":
        "{% if has_children %}References to my children and descendants in this Last Will and Testament shall refer only to {% for child in children %}{{ child.full_name | name }}{% if not loop.last %}, {% endif %}{% endfor %}, the descendants of my aforesaid child or children, any child of mine born or adopted after the date of this Last Will and Testament, and the descendants of any such after-born or after adopted child. A person legally adopted prior to attaining eighteen (18) years of age shall not be differentiated from blood descendants for any purpose, but a person who is adopted after attaining eighteen (18) years of age shall not be deemed a child or descendant for all purposes hereunder.{% endif %}",
    "I devise the residue of my estate to the Trustee then serving under the [TRUST NAME] dated _____________, (hereinafter sometimes referred to as the “Revocable Trust”) executed immediately prior hereto of which I am the Grantor and the present Trustee, to be added to the property held thereunder and administered in accordance with the terms of the Revocable Trust as the same shall exist at the time of my death and not as a trust under this Last Will and Testament.":
        "{% if uses_revocable_trust %}I devise the residue of my estate to the Trustee then serving under the {{ trust.name | name }} dated {{ trust.date }}, (hereinafter sometimes referred to as the “Revocable Trust”) executed immediately prior hereto of which I am the Grantor and the present Trustee, to be added to the property held thereunder and administered in accordance with the terms of the Revocable Trust as the same shall exist at the time of my death and not as a trust under this Last Will and Testament.{% endif %}",
    "I wish to specifically exclude as guardian of the person and the property of any minor child of mine the following individual  for reasons known:":
        "{% if requires_guardian_clause and has_excluded_guardians %}I wish to specifically exclude as guardian of the person and the property of any minor child of mine the following individual(s): {% for person in guardians.excluded %}{{ person.full_name | name }}{% if person.reason %} for {{ person.reason }}{% endif %}{% if not loop.last %}; {% endif %}{% endfor %}.{% endif %}",
    "TO BE DETERMINED":
        "",
    "I declare that the above-named individual(s) shall not be permitted to care for or raise my minor child in the event of my death or disability. Regardless of whether or how much the above-named individual might improve and reform their lives, I do not want them to serve as guardian for my minor child in any circumstances.":
        "{% if requires_guardian_clause and has_excluded_guardians %}I declare that the above-named individual(s) shall not be permitted to care for or raise my minor child in the event of my death or disability. Regardless of whether or how much the above-named individual might improve and reform their lives, I do not want them to serve as guardian for my minor child in any circumstances.{% endif %}",
}


def replace_text_preserve_first_run(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def convert_paragraph(paragraph):
    original = paragraph.text
    normalized = original.strip()
    if normalized in PARAGRAPH_REPLACEMENTS:
        replace_text_preserve_first_run(paragraph, PARAGRAPH_REPLACEMENTS[normalized])
        return
    converted = original
    for before, after in REPLACEMENTS.items():
        converted = converted.replace(before, after)
    if converted != original:
        replace_text_preserve_first_run(paragraph, converted)


def main():
    if not SOURCE.exists():
        raise SystemExit(f"Missing source form: {SOURCE}")
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    for paragraph in doc.paragraphs:
        convert_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    convert_paragraph(paragraph)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
