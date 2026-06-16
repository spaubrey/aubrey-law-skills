#!/usr/bin/env python3
"""Convert the simplified joint revocable trust form into a generator template."""

from pathlib import Path
import re
from tempfile import NamedTemporaryFile
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source-forms" / "trust" / "joint-trust-simplified.docx"
TARGET = ROOT / "templates" / "joint-trust-simplified.docx"


REPLACEMENTS = {
    "[TRUST NAME]": "{{ trust.name | name }}",
    "[CLIENT FULL NAME]": "{{ client.full_name | name }}",
    "[SPOUSE FULL NAME]": "{{ spouse.full_name | name }}",
    "[SIGNING COUNTY]": "{{ execution.display_signing_county | name }}",
    "[Notary Commission]": "{{ execution.display_notary_commission }}",
    "[Seal]": "Seal",
    "ALTERNATE SUCCESSOR TRUSTEE": "{{ fiduciaries.trustee.successor.full_name | name }}",
    "SUCCESSOR TRUSTEE": "{{ fiduciaries.trustee.primary.full_name | name }}",
    "Option 4a": "a majority of our children",
    ",  may name a successor Trustee": ", a majority of our children may name a successor Trustee",
}


PARAGRAPH_REPLACEMENTS = {
    "________________, 20__": "{{ trust.display_date }}",
    "The date of this trust is ________________, 20__.  The parties to this trust are [CLIENT FULL NAME] and [SPOUSE FULL NAME] (the Grantors) and [CLIENT FULL NAME] and [SPOUSE FULL NAME] (collectively, the Trustee).":
        "The date of this trust is {{ trust.display_date }}. The parties to this trust are {{ client.full_name | name }} and {{ spouse.full_name | name }} (the Grantors) and {{ client.full_name | name }} and {{ spouse.full_name | name }} (collectively, the Trustee).",
    "“[CLIENT FULL NAME] and [SPOUSE FULL NAME], Trustees, or their successors in interest, of the [TRUST NAME] dated ________________, 20__, and any amendments thereto.”":
        "“{{ client.full_name | name }} and {{ spouse.full_name | name }}, Trustees, or their successors in interest, of the {{ trust.name | name }} dated {{ trust.display_date }}, and any amendments thereto.”",
    "We have three children.  They are CHILD ONE; CHILD TWO; and CHILD THREE.":
        "{% if has_children %}Our children are {% for child in children %}{{ child.full_name | name }}{% if not loop.last %}, {% endif %}{% endfor %}.{% else %}We have no children living on the date of this trust.{% endif %}",
    "[OPTION – IF CLIENT HAS CHILD FROM PREVIOUS]CHILD ONE is [CLIENT FULL NAME]’s child and not the biological or adopted child of [SPOUSE FULL NAME].  But for the purposes of this trust, CHILD ONE will be considered to be the child of [SPOUSE FULL NAME] and included in references to our children.":
        "{% if trust.blended_family is defined and trust.blended_family.client_children is defined and trust.blended_family.client_children %}{% for child in trust.blended_family.client_children %}{{ child.full_name | name }} is {{ client.full_name | name }}'s child and not the biological or adopted child of {{ spouse.full_name | name }}. But for the purposes of this trust, {{ child.full_name | name }} will be considered to be the child of {{ spouse.full_name | name }} and included in references to our children.{% if not loop.last %} {% endif %}{% endfor %}{% endif %}",
    "[OPTION – IF SPOUSE HAS CHILD FROM PREVIOUS]CHILD THREE is [SPOUSE FULL NAME]’s child and not the biological or adopted child of [CLIENT FULL NAME].  But for the purposes of this trust, CHILD THREE will be considered to be the child of [CLIENT FULL NAME] and included in references to our children.":
        "{% if trust.blended_family is defined and trust.blended_family.spouse_children is defined and trust.blended_family.spouse_children %}{% for child in trust.blended_family.spouse_children %}{{ child.full_name | name }} is {{ spouse.full_name | name }}'s child and not the biological or adopted child of {{ client.full_name | name }}. But for the purposes of this trust, {{ child.full_name | name }} will be considered to be the child of {{ client.full_name | name }} and included in references to our children.{% if not loop.last %} {% endif %}{% endfor %}{% endif %}",
    "We have executed this trust on ________________, 20__.  This trust instrument is effective when signed by us, whether or not now signed by a Trustee.":
        "We have executed this trust on {{ execution.display_doc_date }}. This trust instrument is effective when signed by us, whether or not now signed by a Trustee.",
    "On this day, ________________, 20__, before me, the undersigned notary public, personally appeared [CLIENT FULL NAME], as Grantor and as Trustee, and [SPOUSE FULL NAME], as Grantor and as Trustee, proved to me through satisfactory evidence of identification, which were a government-issued photo identification, to be the persons whose names are signed to the preceding trust instrument, and acknowledged to me that they signed it voluntarily for its stated purposes.":
        "On this day, {{ execution.display_doc_date }}, before me, the undersigned notary public, personally appeared {{ client.full_name | name }}, as Grantor and as Trustee, and {{ spouse.full_name | name }}, as Grantor and as Trustee, proved to me through satisfactory evidence of identification, which were a government-issued photo identification, to be the persons whose names are signed to the preceding trust instrument, and acknowledged to me that they signed it voluntarily for its stated purposes.",
    "Specific Cash Bequests [OPTIONAL]": "{% if has_specific_gifts %}Specific Cash Bequests{% endif %}",
    "[OPTIONAL — Insert specific cash bequests here (e.g., $X to named individual or charity). If no cash bequests, delete this entire Section and this bracketed note before finalizing.]":
        "{% if has_specific_gifts %}The Trustee shall make the following specific gifts: {% for gift in specific_gifts %}{{ gift.item }} to {{ gift.recipient | name }}{% if not loop.last %}; {% endif %}{% endfor %}.{% endif %}",
    "Charitable Gifts [OPTIONAL]": "{% if has_charitable_gifts %}Charitable Gifts{% endif %}",
    "[OPTIONAL — Insert specific charitable gifts here (e.g., gift to named 501(c)(3) organization, charitable remainder interest). If no charitable gifts, delete this entire Section and this bracketed note before finalizing.]":
        "{% if has_charitable_gifts %}The Trustee shall make the following charitable gifts: {% for gift in charitable_gifts %}{{ gift.item }} to {{ gift.recipient | name }}{% if not loop.last %}; {% endif %}{% endfor %}.{% endif %}",
    "Option to Purchase Specific Asset [OPTIONAL]": "",
    "[OPTIONAL — Insert option-to-purchase provisions here (e.g., right of a named beneficiary to purchase a specific asset such as real estate or business interest at appraised value). If no purchase options granted, delete this entire Section and this bracketed note before finalizing.]": "",
    "Pet Trust [OPTIONAL]": "",
    "[OPTIONAL — Insert pet trust provisions here if the Grantors wish to provide for the care of a companion animal after the survivor's death (allocation amount, caregiver designation, distribution of remainder on pet's death). If no pet trust, delete this entire Section and this bracketed note before finalizing.]": "",
    "[Double-Click to Insert Table of Contents]": "",
}


OPTION_MARKER_RE = re.compile(r"\[(?:IF|OPTION)[^\]]*\]")
OPTIONAL_NOTE_RE = re.compile(r"^\[OPTIONAL\s+—")
OOXML_REPLACEMENTS = {
    "[Double-Click to Insert Table of Contents]": "",
}


def replace_text_preserve_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def strip_option_markers(text: str) -> str:
    text = text.replace("[OPTIONAL]", "")
    text = OPTION_MARKER_RE.sub("", text)
    return text.strip() if text.strip() else ""


def convert_paragraph(paragraph) -> None:
    original = paragraph.text
    normalized = original.strip()

    if normalized in PARAGRAPH_REPLACEMENTS:
        replace_text_preserve_first_run(paragraph, PARAGRAPH_REPLACEMENTS[normalized])
        return

    if OPTIONAL_NOTE_RE.match(normalized):
        replace_text_preserve_first_run(paragraph, "")
        return

    converted = original
    for before, after in REPLACEMENTS.items():
        converted = converted.replace(before, after)
    converted = strip_option_markers(converted)
    if converted != original:
        replace_text_preserve_first_run(paragraph, converted)


def clean_ooxml(path: Path) -> None:
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as tmp:
        tmp_path = Path(tmp.name)

    with ZipFile(path) as source, ZipFile(tmp_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", errors="ignore")
                for before, after in OOXML_REPLACEMENTS.items():
                    text = text.replace(before, after)
                data = text.encode("utf-8")
            target.writestr(item, data)
    tmp_path.replace(path)


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"Missing source form: {SOURCE}")
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    for paragraph in doc.paragraphs:
        convert_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    convert_paragraph(paragraph)
    for section in doc.sections:
        story_parts = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for story_part in story_parts:
            for paragraph in story_part.paragraphs:
                convert_paragraph(paragraph)
            for table in story_part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            convert_paragraph(paragraph)
    doc.save(TARGET)
    clean_ooxml(TARGET)


if __name__ == "__main__":
    main()
