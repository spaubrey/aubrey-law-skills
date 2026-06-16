#!/usr/bin/env python3
"""Convert the joint revocable trust form into a generator template."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source-forms" / "trust" / "joint-trust.docx"
TARGET = ROOT / "templates" / "joint-trust.docx"


REPLACEMENTS = {
    "[TRUST NAME]": "{{ trust.name | name }}",
    "[CLIENT FULL NAME]": "{{ client.full_name | name }}",
    "[SPOUSE FULL NAME]": "{{ spouse.full_name | name }}",
    "[SIGNING COUNTY]": "{{ execution.display_signing_county | name }}",
    "[Notary Expiration]": "{{ execution.display_notary_commission }}",
    "[Seal]": "Seal",
    "[OPTIONAL]": "",
}


PARAGRAPH_REPLACEMENTS = {
    "_______________________, 20___": "{{ trust.display_date }}",
    "The date of this trust is _____________________, 20___.  The parties to this trust are [CLIENT FULL NAME] and [SPOUSE FULL NAME] (the Grantors) and [CLIENT FULL NAME] and [SPOUSE FULL NAME] (collectively, our Trustee).":
        "The date of this trust is {{ trust.display_date }}. The parties to this trust are {{ client.full_name | name }} and {{ spouse.full_name | name }} (the Grantors) and {{ client.full_name | name }} and {{ spouse.full_name | name }} (collectively, our Trustee).",
    '"The [TRUST NAME] dated _______________________, 20___."':
        '"The {{ trust.name | name }} dated {{ trust.display_date }}."',
    '"[CLIENT FULL NAME] and [SPOUSE FULL NAME], Trustees, or their successors in interest, of the [TRUST NAME] dated _______________________, 20___, and any amendments thereto."':
        '"{{ client.full_name | name }} and {{ spouse.full_name | name }}, Trustees, or their successors in interest, of the {{ trust.name | name }} dated {{ trust.display_date }}, and any amendments thereto."',
    "We have three children.  They are [CHILD FULL NAME].":
        "{% if has_children %}Our children are {% for child in children %}{{ child.full_name | name }}{% if not loop.last %}, {% endif %}{% endfor %}.{% else %}We have no children living on the date of this trust.{% endif %}",
    "We intentionally have not provided for [DISINHERIT INDIVIDUAL]as a beneficiary in this trust, therefore, for all purposes of this trust, [DISINHERIT INDIVIDUAL]will be treated as having predeceased both of us.":
        "{% if trust.disinherited %}We intentionally have not provided for {% for person in trust.disinherited %}{{ person.full_name | name }}{% if not loop.last %}, {% endif %}{% endfor %} as a beneficiary in this trust; therefore, for all purposes of this trust, each such person will be treated as having predeceased both of us.{% endif %}",
    "TRUSTEE then":
        "{{ fiduciaries.trustee.primary.full_name | name }} then",
    "TRUSTEE then ":
        "{{ fiduciaries.trustee.primary.full_name | name }} then",
    "TRUSTEE 1.":
        "{{ fiduciaries.trustee.successor.full_name | name }}.",
    "We have executed this trust on ___________________, 20___.  This trust instrument is effective when signed by us, whether or not now signed by a Trustee.":
        "We have executed this trust on {{ execution.display_doc_date }}. This trust instrument is effective when signed by us, whether or not now signed by a Trustee.",
    "On this day, _______________________, 20___, before me, the undersigned notary public, personally appeared [CLIENT FULL NAME], as Grantor and as Trustee, and [SPOUSE FULL NAME], as Grantor and as Trustee, proved to me through satisfactory evidence of identification, which were a government-issued photo identification, to be the persons whose names are signed to the preceding trust instrument, and acknowledged to me that they signed it voluntarily for its stated purposes.":
        "On this day, {{ execution.display_doc_date }}, before me, the undersigned notary public, personally appeared {{ client.full_name | name }}, as Grantor and as Trustee, and {{ spouse.full_name | name }}, as Grantor and as Trustee, proved to me through satisfactory evidence of identification, which were a government-issued photo identification, to be the persons whose names are signed to the preceding trust instrument, and acknowledged to me that they signed it voluntarily for its stated purposes.",
}


def replace_text_preserve_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def convert_paragraph(paragraph) -> None:
    original = paragraph.text
    normalized = original.strip()
    if normalized.startswith("[OPTIONAL"):
        replace_text_preserve_first_run(paragraph, "")
        return
    if normalized.startswith("[CLIENT CHILD]"):
        replace_text_preserve_first_run(
            paragraph,
            "{% if trust.blended_family is defined and trust.blended_family.client_children is defined and trust.blended_family.client_children %}{% for child in trust.blended_family.client_children %}{{ child.full_name | name }} is {{ client.full_name | name }}'s child and not the biological or adopted child of {{ spouse.full_name | name }}. But for the purposes of this trust, {{ child.full_name | name }} will be considered to be the child of {{ spouse.full_name | name }} and included in references to our children.{% if not loop.last %} {% endif %}{% endfor %}{% endif %}",
        )
        return
    if normalized.startswith("[SPOUSE CHILD]"):
        replace_text_preserve_first_run(
            paragraph,
            "{% if trust.blended_family is defined and trust.blended_family.spouse_children is defined and trust.blended_family.spouse_children %}{% for child in trust.blended_family.spouse_children %}{{ child.full_name | name }} is {{ spouse.full_name | name }}'s child and not the biological or adopted child of {{ client.full_name | name }}. But for the purposes of this trust, {{ child.full_name | name }} will be considered to be the child of {{ client.full_name | name }} and included in references to our children.{% if not loop.last %} {% endif %}{% endfor %}{% endif %}",
        )
        return
    if normalized in PARAGRAPH_REPLACEMENTS:
        replace_text_preserve_first_run(paragraph, PARAGRAPH_REPLACEMENTS[normalized])
        return
    converted = original
    for before, after in REPLACEMENTS.items():
        converted = converted.replace(before, after)
    if converted != original:
        replace_text_preserve_first_run(paragraph, converted)


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"Missing source form: {SOURCE}")
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    for paragraph in doc.paragraphs:
        convert_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    convert_paragraph(paragraph)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
