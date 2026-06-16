#!/usr/bin/env python3
"""Convert Massachusetts incapacity source forms into generator templates."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "source-forms" / "incapacity" / "v1.1"


OUTPUTS = {
    "Durable_Power_of_Attorney_v1.1.docx": ROOT / "templates" / "durable-power-of-attorney.docx",
    "HCP_v1.1.docx": ROOT / "templates" / "health-care-proxy.docx",
    "Living_Will_v1.1.docx": ROOT / "templates" / "living-will.docx",
    "HIPAA_v1.1.docx": ROOT / "templates" / "hipaa-authorization.docx",
}


REPLACEMENTS = {
    "[CLIENT]": "{{ client.full_name | name }}",
    "[City]": "{{ client.city }}",
    "[Street Address]": "{{ client.street_address }}",
    "[Client DOB]": "{{ client.date_of_birth }}",
    "[SIGNING COUNTY]": "{{ execution.display_signing_county | name }}",
    "[DocDate]": "{{ execution.display_doc_date }}",
    "[Ordinal_DocDate]": "{{ execution.display_ordinal_doc_date }}",
    "[Notary Commission]": "{{ execution.display_notary_commission }}",
    "[INITIAL POA]": "{{ fiduciaries.agent.primary.full_name | name }}",
    "[CO POA]": "{{ fiduciaries.agent.co_primary.full_name | name }}",
    "[SPOUSE]": "{{ spouse.full_name | name }}",
    "[PRIMARY HCP FULL NAME]": "{{ healthcare.agent.primary.full_name | name }}",
    "[Primary HCP Relationship]": "{{ healthcare.agent.primary.relationship }}",
    "[Primary HCP Full Address]": "{{ healthcare.agent.primary.address }}",
    "[Primary HCP Phone]": "{{ healthcare.agent.primary.phone }}",
    "[ALTERNATE 1 HCP FULL NAME]": "{{ healthcare.agent.alternate.full_name | name }}",
    "[Alternate 1 HCP Relationship]": "{{ healthcare.agent.alternate.relationship }}",
    "[Alternate 1 HCP Full Address]": "{{ healthcare.agent.alternate.address }}",
    "[Alternate 1 HCP Phone]": "{{ healthcare.agent.alternate.phone }}",
    "[Client Pronoun]": "{{ client.pronoun_subject }}",
    "[client he/she]": "{{ client.pronoun_subject }}",
    "[Client HisHer]": "{{ client.pronoun_possessive }}",
    "[Spouse HisHer]": "{{ spouse.pronoun_possessive }}",
    "[Remainder of page intentionally left blank]": "Remainder of page intentionally left blank",
}


INLINE_TAGS = {
    "[IF_MARRIED]": "{% if has_spouse %}",
    "[END_IF_MARRIED]": "{% endif %}",
    "[END_IF_AIF_IS_MARRIED]": "{% endif %}",
    "[END_IF_SOLO_AGENT]": "{% endif %}",
    "[END_IF_CO_AGENT]": "{% endif %}",
    "[END_IF_JOINT]": "{% endif %}",
    "[END_IF_SEPARATE]": "{% endif %}",
}


PREFIX_TAGS = {
    "[IF_SOLO_AGENT": "{% if has_solo_poa %}",
    "[IF_CO_AGENT + IF_JOINT": "{% if co_poa_joint %}",
    "[IF_CO_AGENT + IF_SEPARATE": "{% if co_poa_separate %}",
    "[IF_CO_AGENT": "{% if has_co_poa %}",
    "[IF_MARRIED": "{% if has_spouse %}",
    "[IF_AIF_IS_MARRIED": "{% if aif_has_spouse %}",
}


def replace_prefixed_instruction_tags(text: str) -> str:
    changed = True
    while changed:
        changed = False
        for prefix, replacement in PREFIX_TAGS.items():
            start = text.find(prefix)
            if start == -1:
                continue
            end = text.find("]", start)
            if end == -1:
                continue
            text = text[:start] + replacement + text[end + 1 :]
            changed = True
    return text


def convert_text(text: str) -> str:
    converted = replace_prefixed_instruction_tags(text)
    for before, after in INLINE_TAGS.items():
        converted = converted.replace(before, after)
    for before, after in REPLACEMENTS.items():
        converted = converted.replace(before, after)
    return converted


def replace_text_preserve_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def convert_paragraph(paragraph) -> None:
    original = paragraph.text
    converted = convert_text(original)
    if converted != original:
        replace_text_preserve_first_run(paragraph, converted)


def convert_docx(source: Path, target: Path) -> None:
    document = Document(source)
    for paragraph in document.paragraphs:
        convert_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    convert_paragraph(paragraph)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> None:
    for source_name, target in OUTPUTS.items():
        source = SOURCE_DIR / source_name
        if not source.exists():
            raise SystemExit(f"Missing source form: {source}")
        convert_docx(source, target)


if __name__ == "__main__":
    main()
