#!/usr/bin/env python3
"""Create starter Word templates for the estate planning generator."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


TEMPLATES = {
    "will.docx": [
        ("LAST WILL AND TESTAMENT OF {{ client.full_name | name }}", "Title"),
        ("Article I. Family", None),
        ('{{ include_clause("will.family_declaration") }}', None),
        ("Article II. Personal Representative", None),
        ('{{ include_clause("will.executor_appointment") }}', None),
        ("{% if requires_guardian_clause %}Article III. Guardian{% endif %}", None),
        ('{% if requires_guardian_clause %}{{ include_clause("will.guardian_nomination") }}{% endif %}', None),
        ("Article IV. Specific Gifts", None),
        ('{{ include_clause("will.specific_gifts") }}', None),
        ("Article V. Execution", None),
        ('{{ include_clause("execution.signing_block") }}', None),
    ],
    "revocable-trust.docx": [
        ("REVOCABLE LIVING TRUST OF {{ client.full_name | name }}", "Title"),
        ("Article I. Trust Creation", None),
        ("This revocable living trust is created by {{ client.full_name | name }}.", None),
        ("Article II. Successor Trustee", None),
        ('{{ include_clause("trust.trustee_appointment") }}', None),
        ("Article III. Spendthrift Protection", None),
        ('{{ include_clause("trust.spendthrift") }}', None),
    ],
    "durable-power-of-attorney.docx": [
        ("DURABLE POWER OF ATTORNEY OF {{ client.full_name | name }}", "Title"),
        ("Article I. Appointment of Agent", None),
        ('{{ include_clause("poa.agent_appointment") }}', None),
    ],
    "health-care-directive.docx": [
        ("HEALTH CARE DIRECTIVE OF {{ client.full_name | name }}", "Title"),
        ("Article I. Health Care Agent", None),
        ('{{ include_clause("healthcare.agent_appointment") }}', None),
    ],
    "hipaa-authorization.docx": [
        ("HIPAA AUTHORIZATION OF {{ client.full_name | name }}", "Title"),
        ("Article I. Authorization", None),
        ('{{ include_clause("hipaa.authorization") }}', None),
    ],
    "ancillary-execution.docx": [
        ("ANCILLARY EXECUTION DOCUMENTS FOR {{ client.full_name | name }}", "Title"),
        ("Execution Block", None),
        ('{{ include_clause("execution.signing_block") }}', None),
        ("Self-Proving Affidavit", None),
        ('{{ include_clause("execution.self_proving_affidavit", state=state.code) }}', None),
    ],
}


def main() -> None:
    template_dir = ROOT / "templates"
    template_dir.mkdir(parents=True, exist_ok=True)
    for filename, paragraphs in TEMPLATES.items():
        doc = Document()
        for text, style in paragraphs:
            if style:
                doc.add_paragraph(text, style=style)
            else:
                doc.add_paragraph(text)
        doc.save(template_dir / filename)


if __name__ == "__main__":
    main()
