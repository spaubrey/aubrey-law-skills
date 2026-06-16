#!/usr/bin/env python3
"""
render_engagement.py — Fill Aubrey Law engagement agreement templates.

Selects the correct template from assets/ based on plan type and legal plan
status, then replaces text placeholders throughout the document.

Templates (in assets/):
  Engagement_Agreement_C_Trust.docx    — flat-fee, couple, trust plan
  Engagement_Agreement_C-LP_Trust.docx — legal plan, couple, trust plan

Usage:
    python3 render_engagement.py --inputs /tmp/inputs.json --out /tmp/rendered.docx

inputs.json keys:
  client_first_name, client_last_name, client_email  str
  client_phone                                        str (optional)
  couple                                              bool
  spouse_first_name, spouse_last_name, spouse_email  str (if couple)
  spouse_phone                                        str (optional)
  plan_type                                           "trust" | "will"
  trust_type                                          "joint"|"separate"|"individual"
  includes_deed                                       bool
  legal_plan                                          bool
  plan_name                                           str (if legal_plan)
  flat_fee                                            str digits (if not legal_plan)
  cst_addendum                                        bool
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from copy import deepcopy

ASSETS_DIR = Path(__file__).parent.parent / "assets"

# Heading labels where only the label should be bold, not the body text.
# These are the exact Run 0 text strings as they appear in the templates.
HEADING_LABELS = {
    "Identification of Parties:",
    "Legal Fees and Billing Statements:",
}


# ---------------------------------------------------------------------------
# Template selection
# ---------------------------------------------------------------------------

def select_template(inputs: dict) -> Path:
    legal_plan = bool(inputs.get("legal_plan"))
    plan_type = inputs.get("plan_type", "trust").lower()

    if plan_type == "will" and not legal_plan:
        raise SystemExit(
            "ERROR: Will-only flat-fee templates are not yet bundled in this "
            "skill. Draft engagement manually or add will templates to assets/."
        )

    # Will + legal plan uses the LP trust template; Schedule A is adjusted
    # at render time (see apply_will_plan_edits).
    tmpl = (
        ASSETS_DIR / "Engagement_Agreement_C-LP_Trust.docx"
        if legal_plan
        else ASSETS_DIR / "Engagement_Agreement_C_Trust.docx"
    )

    if not tmpl.exists():
        raise SystemExit(f"ERROR: Template not found: {tmpl}")

    return tmpl


# ---------------------------------------------------------------------------
# Placeholder map
# ---------------------------------------------------------------------------

def build_replacements(inputs: dict) -> dict[str, str]:
    client_full = (
        f"{inputs['client_first_name'].strip()} {inputs['client_last_name'].strip()}"
    )

    reps: dict[str, str] = {
        "[Client Full Name]": client_full,
        "[client_full_name]": client_full,
    }

    if inputs.get("couple") and inputs.get("spouse_first_name"):
        spouse_full = (
            f"{inputs['spouse_first_name'].strip()} "
            f"{inputs['spouse_last_name'].strip()}"
        )
        reps["[Spouse Full Name]"] = spouse_full
        reps["[spouse_full_name]"] = spouse_full
    else:
        # Individual client: collapse "X and Y" constructions first, then
        # blank any remaining spouse placeholders.
        reps["[Client Full Name] and [Spouse Full Name]"] = client_full
        reps[" and [Spouse Full Name]"] = ""
        reps["[Spouse Full Name]"] = ""
        reps["[spouse_full_name]"] = ""

    if inputs.get("legal_plan"):
        plan_name = inputs.get("plan_name", "").strip() or "[Legal Plan]"
        reps["[Legal Plan]"] = plan_name
        reps["[legal_plan]"] = plan_name
    else:
        fee_raw = str(inputs.get("flat_fee", "")).strip().lstrip("$").replace(",", "")
        fee_display = f"${int(fee_raw):,}" if fee_raw else "[flat-fee]"
        reps["[flat-fee]"] = fee_display
        reps["[Flat Fee]"] = fee_display

    return reps


# ---------------------------------------------------------------------------
# Run-level paragraph replacement
#
# Strategy: rebuild the full paragraph text, apply replacements, then write
# the result back run-by-run while preserving each run's character formatting.
# For heading paragraphs (Run 0 ends with ':' and is bold), we split the
# result at the ':' boundary so the label stays in Run 0 (bold) and the body
# text moves to a new explicit Run 1 (not bold), preventing the renderer from
# fusing everything into a single bold run.
# ---------------------------------------------------------------------------

def apply_replacements_to_text(text: str, replacements: dict[str, str]) -> str:
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)
    return text


def replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    """Replace placeholders, preserving per-run formatting."""
    runs = para.runs
    if not runs:
        return

    # Build full text with run boundaries marked so we can restore formatting
    full_text = "".join(r.text for r in runs)
    new_text = apply_replacements_to_text(full_text, replacements)

    if new_text == full_text:
        return  # nothing to do

    # Detect heading paragraph: Run 0 is bold and its text ends with ':'
    r0 = runs[0]
    r0_label = apply_replacements_to_text(r0.text, replacements)
    is_heading = bool(r0.bold) and r0_label.rstrip().endswith(":")

    if is_heading:
        # Find the split point: everything up to and including the first ':'
        colon_pos = new_text.find(":")
        if colon_pos != -1:
            label_part = new_text[: colon_pos + 1]
            body_part = new_text[colon_pos + 1 :]
        else:
            label_part = new_text
            body_part = ""

        # Put label into Run 0 (bold preserved from template)
        r0.text = label_part

        # Blank all other runs
        for run in runs[1:]:
            run.text = ""

        # Add a new non-bold run for the body text
        if body_part:
            new_run = para.add_run(body_part)
            new_run.bold = False
            # Copy font name/size from Run 0 so appearance matches
            if r0.font.name:
                new_run.font.name = r0.font.name
            if r0.font.size:
                new_run.font.size = r0.font.size
    else:
        # Non-heading: distribute replaced text back across existing runs
        # proportionally by original run length.
        original_lengths = [len(r.text) for r in runs]
        total_original = sum(original_lengths)

        if total_original == 0:
            runs[0].text = new_text
            return

        # Distribute new_text proportionally
        pos = 0
        for i, run in enumerate(runs):
            if i == len(runs) - 1:
                run.text = new_text[pos:]
            else:
                chars = round(len(new_text) * original_lengths[i] / total_original)
                run.text = new_text[pos: pos + chars]
                pos += chars


def replace_in_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para, replacements)
            for nested in cell.tables:
                replace_in_table(nested, replacements)



# ---------------------------------------------------------------------------
# Structural edits — will plans and individual clients
# ---------------------------------------------------------------------------

def _delete_paragraph(para) -> None:
    para._element.getparent().remove(para._element)


def _delete_table_row(table, row) -> None:
    table._tbl.remove(row._tr)


def apply_will_plan_edits(doc) -> None:
    """Will Plan + legal plan: adjust Schedule A.

    - Remove "Revocable Trust" and "Deed" from covered bullets.
    - Add "Revocable Trust" to the not-covered bullets.
    - Remove Credit Shelter Planning and Deed Recording rows from the
      Optional Services table.
    """
    from copy import deepcopy

    covered_remove = []
    not_covered_anchor = None
    template_bullet = None

    in_covered = False
    in_not_covered = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Your") and "covers the preparation" in t:
            in_covered, in_not_covered = True, False
            continue
        if t.startswith("The plan does not cover"):
            in_covered, in_not_covered = False, True
            continue
        if t == "Optional Services":
            in_covered = in_not_covered = False
            continue
        if in_covered and para.style.name == "List Paragraph":
            if t.startswith("Revocable Trust") or t == "Deed":
                covered_remove.append(para)
        if in_not_covered and para.style.name == "List Paragraph":
            not_covered_anchor = para  # last not-covered bullet
            if template_bullet is None:
                template_bullet = para

    if template_bullet is not None and not_covered_anchor is not None:
        new_el = deepcopy(template_bullet._element)
        not_covered_anchor._element.addnext(new_el)
        # set text on the copied bullet
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_el, not_covered_anchor._parent)
        for i, run in enumerate(new_para.runs):
            run.text = "Revocable Trust" if i == 0 else ""

    for para in covered_remove:
        _delete_paragraph(para)

    # Optional Services table rows
    for table in doc.tables:
        header = " ".join(c.text for c in table.rows[0].cells).lower()
        if "service" in header and "fee" in header:
            for row in list(table.rows[1:]):
                row_text = " ".join(c.text for c in row.cells)
                if "Credit Shelter" in row_text or "Deed Recording" in row_text:
                    _delete_table_row(table, row)


def apply_individual_edits(doc) -> None:
    """Individual (non-couple) client: strip spouse signing blocks and the
    Joint Representation section. Identification clause is handled via
    text replacements in build_replacements()."""
    # 1. Remove Joint Representation paragraphs (from the heading up to,
    #    not including, the Advance Waiver heading).
    deleting = False
    to_delete = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Joint Representation:"):
            deleting = True
        elif deleting and t.startswith("Advance Waiver"):
            deleting = False
        if deleting:
            to_delete.append(para)
    for para in to_delete:
        _delete_paragraph(para)

    # 2. Blank spouse signature cells (any cell containing a signer-2 text
    #    tag or the spouse name placeholder).
    spouse_markers = ("{{signature:2", "{{s:2", "[Spouse Full Name]", "[spouse_full_name]")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(m in cell.text for m in spouse_markers):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""


def render_document(
    template_path: Path, replacements: dict[str, str], out_path: Path,
    inputs: dict | None = None,
) -> None:
    try:
        from docx import Document
    except ImportError:
        raise SystemExit(
            "ERROR: python-docx is not installed. "
            "Run: pip install python-docx --break-system-packages"
        )

    doc = Document(str(template_path))

    inputs = inputs or {}
    if inputs.get("plan_type", "trust").lower() == "will" and inputs.get("legal_plan"):
        apply_will_plan_edits(doc)
    if not inputs.get("couple"):
        apply_individual_edits(doc)

    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        replace_in_table(table, replacements)

    for section in doc.sections:
        for hdr in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            if hdr is None:
                continue
            for para in hdr.paragraphs:
                replace_in_paragraph(para, replacements)
            for table in hdr.tables:
                replace_in_table(table, replacements)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Rendered: {out_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--inputs", required=True, help="Path to inputs JSON")
    p.add_argument("--out", required=True, help="Output .docx path")
    args = p.parse_args()

    inputs = json.loads(Path(args.inputs).read_text())
    template = select_template(inputs)
    replacements = build_replacements(inputs)

    print(f"Template: {template.name}", file=sys.stderr)
    print(f"Replacements: {list(replacements.keys())}", file=sys.stderr)

    render_document(template, replacements, Path(args.out), inputs=inputs)


if __name__ == "__main__":
    main()
